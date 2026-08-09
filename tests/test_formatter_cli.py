from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from conftest import add_activex_ooxml, add_unsafe_ooxml, create_thesis, run_cli


def font_value(paragraph, key: str) -> str | None:
    run = next(run for run in paragraph.runs if run.text.strip())
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get(qn("w:" + key))


def paragraph_by_text(document: Document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def page_number_format(section) -> tuple[str | None, str | None]:
    page_number = section._sectPr.find(qn("w:pgNumType"))
    if page_number is None:
        return None, None
    return page_number.get(qn("w:fmt")), page_number.get(qn("w:start"))


def test_formatter_applies_geometry_font_routing_and_recognized_tables_only(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "fictional-input.docx")
    source_before = source.read_bytes()
    output = tmp_path / "formatted.docx"
    result = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目")
    assert result.returncode == 0, result.stderr
    assert source.read_bytes() == source_before
    document = Document(output)
    for section in document.sections:
        assert round(section.page_width.mm) == 210
        assert round(section.page_height.mm) == 297
        assert round(section.top_margin.mm) == 25
        assert round(section.left_margin.mm) == 30
    chapter = paragraph_by_text(document, "第 1 章 虚构研究")
    heading_one = paragraph_by_text(document, "1.1 已识别一级标题")
    heading_two = paragraph_by_text(document, "1.1.1 已识别二级标题")
    heading_three = paragraph_by_text(document, "1.1.1.1 已识别三级标题")
    body = paragraph_by_text(document, "这是供自动化测试使用的虚构正文。")
    for paragraph, expected_size in ((chapter, 18), (heading_one, 15), (heading_two, 14), (heading_three, 12)):
        assert round(next(run for run in paragraph.runs if run.text.strip()).font.size.pt) == expected_size
        assert font_value(paragraph, "eastAsia") == "黑体"
        assert font_value(paragraph, "ascii") == "Times New Roman"
    assert round(next(run for run in body.runs if run.text.strip()).font.size.pt) == 12
    assert font_value(body, "eastAsia") == "宋体"
    recognized_borders = document.tables[0].cell(0, 0)._tc.tcPr.first_child_found_in("w:tcBorders")
    unrecognized_borders = document.tables[1].cell(0, 0)._tc.tcPr.first_child_found_in("w:tcBorders")
    assert recognized_borders is not None
    assert unrecognized_borders is None
    assert document.tables[1].cell(0, 0).text == "KEEP_UNRECOGNIZED_TABLE"


def test_formatter_uses_roman_front_matter_arabic_main_and_safe_header_border(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "numbering-input.docx")
    output = tmp_path / "numbering-output.docx"
    result = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目")
    assert result.returncode == 0, result.stderr
    document = Document(output)
    assert [page_number_format(section) for section in document.sections] == [("upperRoman", "1"), ("decimal", "1")]
    border = document.sections[0].header.paragraphs[0]._p.pPr.find(qn("w:pBdr")).find(qn("w:bottom"))
    assert border.get(qn("w:val")) == "thickThinSmallGap"
    declaration_index = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == "学位论文原创性声明")
    authorization_index = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == "学位论文版权使用授权书")
    assert declaration_index + 1 == authorization_index
    assert all(paragraph._p.pPr is None or paragraph._p.pPr.sectPr is None for paragraph in document.paragraphs[:authorization_index])


def test_formatter_never_overwrites_input_or_existing_output_without_explicit_permission(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "guarded-input.docx")
    source_before = source.read_bytes()
    output = tmp_path / "existing-output.docx"
    output.write_bytes(b"DO_NOT_OVERWRITE")
    blocked = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目")
    assert blocked.returncode != 0
    assert output.read_bytes() == b"DO_NOT_OVERWRITE"
    same_path = run_cli("scripts.format_thesis", str(source), str(source), "--title", "虚构测试题目", "--overwrite")
    assert same_path.returncode != 0
    assert source.read_bytes() == source_before
    allowed = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目", "--overwrite")
    assert allowed.returncode == 0, allowed.stderr
    assert output.read_bytes() != b"DO_NOT_OVERWRITE"


def test_formatter_guards_report_collisions_and_supports_private_title_file(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "report-guard-input.docx")
    output = tmp_path / "report-guard-output.docx"
    report = tmp_path / "existing-report.json"
    report.write_text("DO_NOT_OVERWRITE", encoding="utf-8")
    blocked = run_cli(
        "scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目", "--json-report", str(report)
    )
    assert blocked.returncode != 0
    assert not output.exists()
    assert report.read_text(encoding="utf-8") == "DO_NOT_OVERWRITE"
    colliding = run_cli(
        "scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目", "--json-report", str(output)
    )
    assert colliding.returncode != 0
    title_file = tmp_path / "title.txt"
    title_file.write_text("虚构测试题目\n", encoding="utf-8")
    allowed = run_cli("scripts.format_thesis", str(source), str(output), "--title-file", str(title_file))
    assert allowed.returncode == 0, allowed.stderr


def test_formatter_rejects_external_relationships_and_automatic_field_updates(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "unsafe-input.docx")
    add_unsafe_ooxml(source)
    output = tmp_path / "unsafe-output.docx"
    result = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目")
    assert result.returncode != 0
    assert not output.exists()


def test_formatter_rejects_activex_controls(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "activex-input.docx")
    add_activex_ooxml(source)
    output = tmp_path / "activex-output.docx"
    result = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目")
    assert result.returncode != 0
    assert not output.exists()


def test_formatter_preserves_nonempty_inherited_headers(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "existing-header-input.docx")
    document = Document(source)
    document.sections[0].header.paragraphs[0].text = "合成既有页眉"
    document.save(source)
    output = tmp_path / "existing-header-output.docx"
    result = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目")
    assert result.returncode == 0, result.stderr
    formatted = Document(output)
    assert [section.header.paragraphs[0].text for section in formatted.sections] == [
        "合成既有页眉",
        "合成既有页眉",
    ]


def test_formatter_preserves_footnote_policy_unless_explicitly_requested(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "footnote-policy-input.docx")
    document = Document(source)
    settings = document.settings._element
    footnote_pr = settings.find(qn("w:footnotePr"))
    if footnote_pr is None:
        footnote_pr = OxmlElement("w:footnotePr")
        settings.append(footnote_pr)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "lowerLetter")
    restart = OxmlElement("w:numRestart")
    restart.set(qn("w:val"), "continuous")
    footnote_pr.extend((num_fmt, restart))
    document.save(source)

    preserved_output = tmp_path / "footnote-preserved.docx"
    preserved = run_cli(
        "scripts.format_thesis", str(source), str(preserved_output), "--title", "虚构测试题目"
    )
    assert preserved.returncode == 0, preserved.stderr
    preserved_pr = Document(preserved_output).settings._element.find(qn("w:footnotePr"))
    assert preserved_pr.find(qn("w:numFmt")).get(qn("w:val")) == "lowerLetter"
    assert preserved_pr.find(qn("w:numRestart")).get(qn("w:val")) == "continuous"

    configured_output = tmp_path / "footnote-configured.docx"
    configured = run_cli(
        "scripts.format_thesis",
        str(source),
        str(configured_output),
        "--title",
        "虚构测试题目",
        "--humanities-footnotes",
    )
    assert configured.returncode == 0, configured.stderr
    configured_pr = Document(configured_output).settings._element.find(qn("w:footnotePr"))
    assert configured_pr.find(qn("w:numFmt")).get(qn("w:val")) == "decimalEnclosedCircle"
    assert configured_pr.find(qn("w:numRestart")).get(qn("w:val")) == "eachPage"


def test_formatter_returns_nonzero_when_postformat_validation_has_errors(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "postvalidation-input.docx")
    document = Document(source)
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("第 1 章"):
            paragraph.clear()
            paragraph.add_run("第 1 章 虚构研究 [1]")
    document.save(source)
    output = tmp_path / "postvalidation-output.docx"
    result = run_cli("scripts.format_thesis", str(source), str(output), "--title", "虚构测试题目")
    assert result.returncode != 0
    assert output.exists()
