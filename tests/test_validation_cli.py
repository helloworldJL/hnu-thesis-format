from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Mm

from conftest import PRIVATE_MARKER, create_thesis, run_cli


def test_validation_reports_page_geometry_without_path_or_excerpt_leaks(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "fictional-input.docx")
    document = Document(source)
    document.sections[0].left_margin = Mm(34)
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("第 1 章"):
            paragraph.clear()
            paragraph.add_run("第 1 章 " + PRIVATE_MARKER)
    document.save(source)
    first_json = tmp_path / "first.json"
    first_html = tmp_path / "first.html"
    second_json = tmp_path / "second.json"
    second_html = tmp_path / "second.html"
    first = run_cli("scripts.validate", str(source), "--json-report", str(first_json), "--html-report", str(first_html))
    second = run_cli("scripts.validate", str(source), "--json-report", str(second_json), "--html-report", str(second_html))
    assert first.returncode == 1
    assert second.returncode == 1
    assert first_json.read_bytes() == second_json.read_bytes()
    assert first_html.read_bytes() == second_html.read_bytes()
    payload = json.loads(first_json.read_text(encoding="utf-8"))
    assert any(issue["category"] == "页面" for issue in payload["issues"])
    report_text = first_json.read_text(encoding="utf-8") + first_html.read_text(encoding="utf-8")
    assert str(source) not in report_text
    assert PRIVATE_MARKER not in report_text
    assert "/" + "Users/" not in report_text
    assert all("Paragraph " in issue["location"] or issue["location"] != "" for issue in payload["issues"])


def test_validation_errors_when_no_reference_items_are_recognized(tmp_path: Path) -> None:
    source = create_thesis(tmp_path / "zero-references.docx")
    document = Document(source)
    for paragraph in list(document.paragraphs):
        if paragraph.text.startswith("[1]"):
            paragraph._element.getparent().remove(paragraph._element)
    document.save(source)
    report = tmp_path / "zero-references.json"
    result = run_cli("scripts.validate", str(source), "--json-report", str(report))
    assert result.returncode == 1
    issues = json.loads(report.read_text(encoding="utf-8"))["issues"]
    assert any(
        issue["severity"] == "error" and issue["rule_id"] == "REFERENCE_ITEMS_MISSING"
        for issue in issues
    )
