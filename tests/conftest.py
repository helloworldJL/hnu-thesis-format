from __future__ import annotations

import subprocess
import sys
import zipfile
import os
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
SKILL_ROOT = REPOSITORY_ROOT / "hnu-thesis-format"
PRIVATE_MARKER = "FICTIONAL_PRIVATE_TEXT_DO_NOT_DISCLOSE"


def run_cli(module: str, *arguments: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, "-m", module, *arguments],
        cwd=SKILL_ROOT,
        check=False,
        capture_output=True,
        env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"},
        text=True,
    )


def set_run_font(paragraph, text: str, size: float, east_asia: str, latin: str) -> None:
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = latin
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_a4_geometry(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(30)
        section.right_margin = Mm(30)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


def add_paragraph(document: Document, text: str, size: float = 9, font: str = "Arial") -> None:
    paragraph = document.add_paragraph()
    set_run_font(paragraph, text, size, font, font)


def create_thesis(path: Path) -> Path:
    document = Document()
    set_a4_geometry(document)
    for text in (
        "学位论文原创性声明",
        "学位论文版权使用授权书",
        "摘 要",
        PRIVATE_MARKER,
        "关键词：虚构；测试；排版",
        "Abstract",
        "A fictional abstract for deterministic testing.",
        "Key Words: fictional; testing; formatting",
        "目 录",
    ):
        add_paragraph(document, text)
    document.add_section(WD_SECTION.NEW_PAGE)
    set_a4_geometry(document)
    for text in (
        "第 1 章 虚构研究",
        "1.1 已识别一级标题",
        "1.1.1 已识别二级标题",
        "1.1.1.1 已识别三级标题",
        "这是供自动化测试使用的虚构正文。",
        "表 1.1 已识别表题",
    ):
        add_paragraph(document, text)
    recognized_table = document.add_table(rows=2, cols=2)
    recognized_table.cell(0, 0).text = "甲"
    recognized_table.cell(0, 1).text = "乙"
    recognized_table.cell(1, 0).text = "1"
    recognized_table.cell(1, 1).text = "2"
    add_paragraph(document, "Plain synthetic table")
    unrecognized_table = document.add_table(rows=1, cols=1)
    unrecognized_table.cell(0, 0).text = "KEEP_UNRECOGNIZED_TABLE"
    for text in (
        "参考文献",
        "[1] Fictional, A. A synthetic reference.",
        "攻读学位期间的学术或实践成果",
        "致 谢",
        "答辩委员会名单",
    ):
        add_paragraph(document, text)
    document.save(path)
    return path


def create_sectioned_thesis(path: Path) -> Path:
    document = Document()
    set_a4_geometry(document)
    add_paragraph(document, "摘 要")
    document.add_section(WD_SECTION.NEW_PAGE)
    set_a4_geometry(document)
    add_paragraph(document, "Abstract")
    document.add_section(WD_SECTION.NEW_PAGE)
    set_a4_geometry(document)
    add_paragraph(document, "目 录")
    document.add_section(WD_SECTION.NEW_PAGE)
    set_a4_geometry(document)
    add_paragraph(document, "第 1 章 虚构正文")
    document.save(path)
    return path


def add_unsafe_ooxml(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source:
        members = {member.filename: source.read(member.filename) for member in source.infolist()}
    relationship_name = "word/_rels/document.xml.rels"
    relationships = members[relationship_name].decode("utf-8")
    relationships = relationships.replace(
        "</Relationships>",
        '<Relationship Id="rIdSyntheticExternal" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate" '
        'Target="https://example.invalid/synthetic" TargetMode="External"/></Relationships>',
    )
    members[relationship_name] = relationships.encode("utf-8")
    settings_name = "word/settings.xml"
    settings = members[settings_name].decode("utf-8")
    members[settings_name] = settings.replace(
        "</w:settings>", '<w:updateFields w:val="true"/></w:settings>'
    ).encode("utf-8")
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for name, content in members.items():
            target.writestr(name, content)


def add_activex_ooxml(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source:
        members = {member.filename: source.read(member.filename) for member in source.infolist()}
    relationship_name = "word/_rels/document.xml.rels"
    relationships = members[relationship_name].decode("utf-8")
    relationships = relationships.replace(
        "</Relationships>",
        '<Relationship Id="rIdSyntheticControl" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/control" '
        'Target="activeX/activeX1.xml"/></Relationships>',
    )
    members[relationship_name] = relationships.encode("utf-8")
    members["word/activeX/activeX1.xml"] = b'<?xml version="1.0" encoding="UTF-8"?><ax:ocx xmlns:ax="urn:synthetic"/>'
    members["word/activeX/activeX1.bin"] = b"SYNTHETIC_ACTIVEX"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for name, content in members.items():
            target.writestr(name, content)


@pytest.fixture
def thesis_path(tmp_path: Path) -> Path:
    return create_thesis(tmp_path / "fictional-input.docx")
