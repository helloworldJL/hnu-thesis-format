from __future__ import annotations

from typing import Optional

from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from . import constants as C

ALIGNMENT = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
             "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}


def _is_field_run(run) -> bool:
    return bool(run._r.xpath(".//w:fldChar | .//w:instrText"))


def set_run_font(run, font_cn: str, font_latin: str, pt: float) -> None:
    if _is_field_run(run):
        return
    run.font.name = font_latin
    run.font.size = Pt(pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key, value in (("eastAsia", font_cn), ("ascii", font_latin),
                       ("hAnsi", font_latin), ("cs", font_latin)):
        rfonts.set(qn("w:" + key), value)


def apply_paragraph_spec(paragraph, spec) -> None:
    fmt = paragraph.paragraph_format
    paragraph.alignment = ALIGNMENT.get(spec.get("alignment"))
    for key, attr in (("space_before_pt", "space_before"), ("space_after_pt", "space_after"),
                      ("line_spacing_pt", "line_spacing"), ("first_line_indent_pt", "first_line_indent")):
        if key in spec:
            setattr(fmt, attr, Pt(spec[key]))
    if "hanging_indent_pt" in spec:
        fmt.first_line_indent = Pt(-spec["hanging_indent_pt"])
        fmt.left_indent = Pt(spec["hanging_indent_pt"])
    if spec.get("keep_with_next"):
        fmt.keep_with_next = True
    if spec.get("page_break_before"):
        fmt.page_break_before = True
    for run in paragraph.runs:
        set_run_font(run, spec["font_cn"], spec["font_latin"], spec["pt"])


def setup_section_page(section) -> None:
    landscape = bool(section.page_width is not None and section.page_height is not None
                     and section.page_width > section.page_height)
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Mm(C.PAGE["height_mm"] if landscape else C.PAGE["width_mm"])
    section.page_height = Mm(C.PAGE["width_mm"] if landscape else C.PAGE["height_mm"])
    section.top_margin = Mm(C.PAGE["margin_top_mm"])
    section.bottom_margin = Mm(C.PAGE["margin_bottom_mm"])
    section.left_margin = Mm(C.PAGE["margin_left_mm"])
    section.right_margin = Mm(C.PAGE["margin_right_mm"])
    section.header_distance = Mm(C.PAGE["header_distance_mm"])
    section.footer_distance = Mm(C.PAGE["footer_distance_mm"])


def ensure_even_odd_headers(document) -> None:
    settings = document.settings._element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def _container_is_empty(container) -> bool:
    if container.tables:
        return False
    return not any(
        p.text.strip()
        or p._p.xpath(
            ".//w:t | .//w:fldChar | .//w:drawing | .//w:pict | .//w:object | .//w:sdt"
        )
        for p in container.paragraphs
    )


def _add_bottom_border(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    for key, value in (("val", C.HEADER["border_val"]),
                       ("sz", C.HEADER["border_size_eighths"]), ("space", 1), ("color", "auto")):
        bottom.set(qn("w:" + key), str(value))


def write_header_if_empty(header, text: str) -> bool:
    if not _container_is_empty(header):
        return False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, C.HEADER["font_cn"], C.HEADER["font_latin"], C.HEADER["pt"])
    _add_bottom_border(paragraph)
    return True


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def write_footer_if_empty(footer) -> bool:
    if not _container_is_empty(footer):
        return False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(paragraph)
    return True


def set_page_number_format(section, fmt: str, start: Optional[int] = None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if start is None:
        pg_num.attrib.pop(qn("w:start"), None)
    else:
        pg_num.set(qn("w:start"), str(start))


def configure_footnote_settings(document) -> None:
    settings = document.settings._element
    footnote_pr = settings.find(qn("w:footnotePr"))
    if footnote_pr is None:
        footnote_pr = OxmlElement("w:footnotePr")
        settings.append(footnote_pr)
    for tag in ("w:numFmt", "w:numRestart"):
        old = footnote_pr.find(qn(tag))
        if old is not None:
            footnote_pr.remove(old)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimalEnclosedCircle")
    restart = OxmlElement("w:numRestart")
    restart.set(qn("w:val"), "eachPage")
    footnote_pr.extend((num_fmt, restart))


def set_cell_border(cell, **edges) -> None:
    borders = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        cell._tc.get_or_add_tcPr().append(borders)
    for edge, attrs in edges.items():
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn("w:" + key), str(value))


def apply_three_line_table(table) -> None:
    rows = table.rows
    if not rows:
        return
    nil = {"val": "nil"}
    for row in rows:
        for cell in row.cells:
            set_cell_border(cell, left=nil, right=nil, top=nil, bottom=nil)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, C.FONT_SONGTI, C.FONT_LATIN, 10.5)
    thick = {"val": "single", "sz": C.TABLE["top_bottom_border_size_eighths"], "color": "auto"}
    thin = {"val": "single", "sz": C.TABLE["middle_border_size_eighths"], "color": "auto"}
    for cell in rows[0].cells:
        set_cell_border(cell, top=thick, bottom=thin, left=nil, right=nil)
    for cell in rows[-1].cells:
        set_cell_border(cell, bottom=thick, left=nil, right=nil)
