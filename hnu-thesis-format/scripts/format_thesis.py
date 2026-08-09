from __future__ import annotations

import argparse
import json
import re
import tempfile
from dataclasses import asdict
from pathlib import Path

from docx.oxml.ns import qn

from . import constants as C
from .report_html import write_html_report
from .safety import (UnsafeDocumentError, atomic_write_text, load_docx,
                     preflight_docx, publish_atomic, validate_output_paths)
from .utils import (apply_paragraph_spec, apply_three_line_table, configure_footnote_settings,
                    ensure_even_odd_headers, set_page_number_format, setup_section_page,
                    write_footer_if_empty, write_header_if_empty)
from .validate import (classify_paragraph, collect_stats, find_landmarks,
                       validate_document)


def _in_reference_range(text, state):
    if re.match(C.REGEX["reference_heading"], text):
        return True
    if state and any(re.match(C.REGEX[key], text) for key in ("appendix", "achievements", "acknowledgements", "defense")):
        return False
    return state


def format_paragraphs(doc):
    counts = {}
    landmarks = find_landmarks(doc)
    main_index = landmarks.get("main", (len(doc.paragraphs), 0))[0]
    toc_index = landmarks.get("toc", (-1, 0))[0]
    in_references = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        kind = classify_paragraph(text)
        in_toc = toc_index >= 0 and toc_index < index < main_index
        style_name = paragraph.style.name.lower() if paragraph.style is not None else ""
        if in_toc:
            if style_name in {"toc 1", "目录 1"}:
                apply_paragraph_spec(paragraph, C.STYLE_SPECS["toc_level1"])
                counts["toc_level1"] = counts.get("toc_level1", 0) + 1
            continue
        in_references = _in_reference_range(text, in_references)
        if in_references and re.match(C.REGEX["reference_item"], text):
            kind = "reference"
        if kind is None and index >= main_index and not in_references:
            kind = "body"
        if index < main_index and kind not in {"abstract_title_zh", "abstract_title_en", "keyword_zh", "keyword_en", "toc_title"}:
            continue
        if kind not in C.STYLE_SPECS:
            continue
        apply_paragraph_spec(paragraph, C.STYLE_SPECS[kind])
        counts[kind] = counts.get(kind, 0) + 1
    return counts


def _paragraph_text(element):
    return "".join(element.xpath(".//w:t/text()"))


def _recognized_table_caption(table):
    sibling = table._tbl.getprevious()
    while sibling is not None:
        if sibling.tag == qn("w:p"):
            text = _paragraph_text(sibling).strip()
            if text:
                return bool(re.match(C.REGEX["caption_table"], text))
        elif sibling.tag == qn("w:tbl"):
            return False
        sibling = sibling.getprevious()
    return False


def format_tables(doc):
    count = 0
    for table in doc.tables:
        if _recognized_table_caption(table):
            apply_three_line_table(table)
            count += 1
    return count


def _write_section_headers(section, odd_text, even_text):
    changed = 0
    preserved = 0
    for header, text in ((section.header, odd_text), (section.even_page_header, even_text)):
        if _container_has_content(header):
            preserved += 1
            continue
        if header.is_linked_to_previous:
            header.is_linked_to_previous = False
        changed += int(write_header_if_empty(header, text))
    for footer in (section.footer, section.even_page_footer):
        if _container_has_content(footer):
            preserved += 1
            continue
        if footer.is_linked_to_previous:
            footer.is_linked_to_previous = False
        changed += int(write_footer_if_empty(footer))
    return changed, preserved


def _container_has_content(container):
    if container.tables:
        return True
    return any(
        paragraph.text.strip()
        or paragraph._p.xpath(
            ".//w:t | .//w:fldChar | .//w:drawing | .//w:pict | .//w:object | .//w:sdt"
        )
        for paragraph in container.paragraphs
    )


def setup_headers_and_numbering(doc, thesis_type, title):
    landmarks = find_landmarks(doc)
    if "abstract" not in landmarks or "main" not in landmarks:
        return {"status": "not_applied_landmark_missing", "sections": 0, "containers_written": 0}
    abstract_section = landmarks["abstract"][1]
    main_section = landmarks["main"][1]
    if main_section <= abstract_section:
        return {"status": "not_applied_section_boundary_ambiguous", "sections": 0, "containers_written": 0}
    ensure_even_odd_headers(doc)
    config = C.THESIS_TYPES[thesis_type]
    written = 0
    preserved = 0
    touched = 0
    for index in range(abstract_section, len(doc.sections)):
        section = doc.sections[index]
        fmt = "upperRoman" if index < main_section else "decimal"
        set_page_number_format(section, fmt, 1 if index in (abstract_section, main_section) else None)
        section_written, section_preserved = _write_section_headers(
            section, config["header_odd"], title
        )
        written += section_written
        preserved += section_preserved
        touched += 1
    return {"status": "applied", "sections": touched, "containers_written": written,
            "containers_preserved": preserved,
            "roman_start_section": abstract_section + 1, "arabic_start_section": main_section + 1}


def format_document(doc, thesis_type, title, humanities_footnotes=False):
    if thesis_type not in C.THESIS_TYPES:
        raise ValueError("Unsupported thesis type")
    stats = {"sections_page_setup": 0}
    for section in doc.sections:
        setup_section_page(section)
        stats["sections_page_setup"] += 1
    stats["paragraphs_by_style"] = format_paragraphs(doc)
    stats["tables_formatted"] = format_tables(doc)
    stats["numbering"] = setup_headers_and_numbering(doc, thesis_type, title)
    stats["footnote_settings"] = "preserved"
    if humanities_footnotes:
        configure_footnote_settings(doc)
        stats["footnote_settings"] = "circled_arabic_each_page"
    return stats


def _save_atomic(doc, output, overwrite):
    output.parent.mkdir(parents=True, exist_ok=True)
    handle = tempfile.NamedTemporaryFile(prefix=".hnu-format-", suffix=".docx", dir=str(output.parent), delete=False)
    temporary = Path(handle.name)
    handle.close()
    try:
        doc.save(str(temporary))
        preflight_docx(temporary)
        publish_atomic(temporary, output, overwrite)
    finally:
        if temporary.exists():
            temporary.unlink()


def _read_title(args):
    if args.title_file:
        try:
            title = Path(args.title_file).read_text(encoding="utf-8").strip()
        except UnicodeError as exc:
            raise ValueError("Title file must contain valid UTF-8 text.") from exc
    else:
        title = args.title.strip()
    if not title:
        raise ValueError("Thesis title must not be empty.")
    if "\n" in title or "\r" in title:
        raise ValueError("Thesis title must be a single line.")
    if len(title) > 200:
        raise ValueError("Thesis title must not exceed 200 characters.")
    return title


def main(argv=None):
    parser = argparse.ArgumentParser(description="Format a copy using conservative HNU 2025 academic doctoral dissertation rules.")
    parser.add_argument("input_docx"); parser.add_argument("output_docx")
    parser.add_argument("--thesis-type", default=C.THESIS_TYPE, choices=sorted(C.THESIS_TYPES))
    title_group = parser.add_mutually_exclusive_group(required=True)
    title_group.add_argument("--title", help="Thesis title; visible in process listings. Prefer --title-file for privacy.")
    title_group.add_argument("--title-file", help="Read the thesis title from a UTF-8 text file and strip surrounding whitespace.")
    parser.add_argument("--overwrite", action="store_true", help="Replace only explicitly requested output/report files.")
    parser.add_argument("--html-report"); parser.add_argument("--json-report")
    parser.add_argument("--report-language", choices=("zh", "en", "bilingual"), default="bilingual")
    parser.add_argument("--include-sensitive-details", action="store_true")
    parser.add_argument(
        "--humanities-footnotes",
        action="store_true",
        help="Apply circled-Arabic footnotes restarted on each page after confirming the rule applies.",
    )
    args = parser.parse_args(argv)
    input_path = Path(args.input_docx)
    output_path = Path(args.output_docx)
    try:
        title = _read_title(args)
        validate_output_paths((input_path, args.title_file),
                              (output_path, args.html_report, args.json_report), args.overwrite)
        doc = load_docx(input_path)
        change_stats = format_document(
            doc, args.thesis_type, title, humanities_footnotes=args.humanities_footnotes
        )
        _save_atomic(doc, output_path, args.overwrite)
        output_doc = load_docx(output_path)
        issues = validate_document(output_doc, args.thesis_type, args.include_sensitive_details)
        stats = collect_stats(output_doc, args.thesis_type)
        stats.update(change_stats)
        if args.json_report:
            payload = {"schema_version": 1, "thesis_type": args.thesis_type, "sources": list(C.SOURCES),
                       "issues": [asdict(issue) for issue in issues], "stats": stats}
            atomic_write_text(args.json_report, json.dumps(payload, ensure_ascii=False, indent=2), args.overwrite)
        if args.html_report:
            write_html_report(issues, args.html_report, overwrite=args.overwrite,
                              thesis_type=args.thesis_type, stats=stats, language=args.report_language)
    except (UnsafeDocumentError, OSError, ValueError) as exc:
        parser.error(str(exc))
    counts = {severity: sum(issue.severity == severity for issue in issues)
              for severity in ("error", "warning", "info")}
    print("Formatting completed: errors={error}, warnings={warning}, info={info}.".format(**counts))
    return 1 if counts["error"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
